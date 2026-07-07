"""
DocxBuilder: renders the final .docx from a DocumentContext + Plan.

- Title page: doc type (humanized) + generation date (from get_current_date
  tool output if available, else today()).
- Styled headings per section (using python-docx's built-in Heading styles).
- "Assumptions Made" section rendered whenever plan.assumptions is non-empty.
- Section ordering follows the order tasks appear in the plan, then any
  validator-inserted placeholders at the end.
"""
import os
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.core.config import settings
from app.models.schemas import DocumentContext, Plan


def _humanize(doc_type: str) -> str:
    return doc_type.replace("_", " ").title()


def _find_generated_date(context: DocumentContext) -> str | None:
    """
    Tool outputs are stored either flat ({'date': ...}) for a single tool call,
    or nested ({'tool_name': {...}}) when a task's LLM call invoked multiple
    tools. Check both shapes for a 'date' key.
    """
    for output in context.tool_outputs.values():
        if not isinstance(output, dict):
            continue
        if "date" in output:
            return output["date"]
        for nested in output.values():
            if isinstance(nested, dict) and "date" in nested:
                return nested["date"]
    return None


def _add_title_page(doc: Document, plan: Plan, context: DocumentContext) -> None:
    gen_date = _find_generated_date(context) or date.today().isoformat()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(_humanize(plan.doc_type))
    run.bold = True
    run.font.size = Pt(28)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Generated {gen_date}")
    sub_run.italic = True
    sub_run.font.size = Pt(12)

    doc.add_page_break()


def _add_assumptions_section(doc: Document, plan: Plan) -> None:
    if not plan.assumptions:
        return
    doc.add_heading("Assumptions Made", level=1)
    for assumption in plan.assumptions:
        doc.add_paragraph(assumption, style="List Bullet")


def _add_section_body(doc: Document, text: str) -> None:
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)


def build_docx(plan: Plan, context: DocumentContext) -> str:
    """Builds the .docx file and returns its path."""
    doc = Document()
    _add_title_page(doc, plan, context)
    _add_assumptions_section(doc, plan)

    # Generated sections, in plan task order
    for task in plan.tasks:
        if task.type != "generate":
            continue
        content = context.sections.get(task.id)
        if not content:
            continue
        doc.add_heading(task.description, level=1)
        _add_section_body(doc, content)

    # Validator-inserted placeholders (keys prefixed accordingly), appended last
    placeholder_items = {k: v for k, v in context.sections.items() if k.startswith("_validator_fill_")}
    if placeholder_items:
        doc.add_heading("Additional Notes", level=1)
        for content in placeholder_items.values():
            doc.add_paragraph(content)

    os.makedirs(settings.output_dir, exist_ok=True)
    safe_type = plan.doc_type.replace(" ", "_").lower()
    timestamp = date.today().isoformat()
    filename = f"{safe_type}_{timestamp}_{abs(hash(context.original_request)) % 10000}.docx"
    path = os.path.join(settings.output_dir, filename)
    doc.save(path)
    return path
