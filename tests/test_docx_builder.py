"""
Test: DocxBuilder produces a non-empty, valid .docx file from a
DocumentContext + Plan.
"""
import os

from docx import Document

from app.documents.docx_builder import build_docx
from app.models.schemas import DocumentContext, Plan, Task


def test_docx_builder_produces_valid_nonempty_docx(tmp_path, monkeypatch):
    monkeypatch.setattr("app.documents.docx_builder.settings.output_dir", str(tmp_path))

    plan = Plan(
        doc_type="proposal",
        assumptions=["Assumed a 3-month timeline since none was specified."],
        tasks=[
            Task(id="t1", description="Look up sections", type="tool", depends_on=[], tool_name="lookup_template"),
            Task(id="t2", description="Executive Summary", type="generate", depends_on=["t1"]),
        ],
    )
    context = DocumentContext(
        doc_type="proposal",
        original_request="Write a proposal for a new client engagement",
        assumptions=plan.assumptions,
    )
    context.set_tool_output("t1", {"doc_type": "proposal", "sections": ["Executive Summary", "Budget"]})
    context.set_section("t2", "This proposal outlines our approach.\n- Key point one\n- Key point two")

    path = build_docx(plan, context)

    assert os.path.exists(path)
    assert path.endswith(".docx")
    assert os.path.getsize(path) > 0

    doc = Document(path)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert len(doc.paragraphs) > 0
    assert "Executive Summary" in all_text
    assert "Assumptions Made" in all_text
